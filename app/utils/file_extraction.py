"""
File content extraction utilities for different file formats.
Supports PDF, DOCX, TXT and other text-based file formats.
"""

import io
import logging
from typing import Optional, Tuple
import tempfile
import os

import fitz

logger = logging.getLogger(__name__)


class FileContentExtractor:
	"""Utility class for extracting text content from various file formats"""

	@staticmethod
	def extract_text_content(file_content: bytes, file_type: str, file_name: str) -> Tuple[Optional[str], Optional[str]]:
		"""
		Extract text content from file bytes based on file type

		Args:
		    file_content: Raw file content as bytes
		    file_type: MIME type of the file
		    file_name: Original filename (for extension fallback)

		Returns:
		    Tuple of (extracted_text, error_message)
		"""
		try:
			# Normalize MIME type
			file_type_lower = file_type.lower()

			# Handle different file types
			if file_type_lower == 'text/plain' or file_name.lower().endswith('.txt'):
				return FileContentExtractor._extract_from_txt(file_content)

			elif file_type_lower == 'application/pdf' or file_name.lower().endswith('.pdf'):
				return FileContentExtractor._extract_from_pdf(file_content)

			elif file_type_lower.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml') or file_name.lower().endswith('.docx'):
				return FileContentExtractor._extract_from_docx(file_content)

			elif file_type_lower == 'application/msword' or file_name.lower().endswith('.doc'):
				return FileContentExtractor._extract_from_doc(file_content)

			elif file_type_lower == 'text/markdown' or file_name.lower().endswith(('.md', '.markdown')):
				return FileContentExtractor._extract_from_txt(file_content)

			elif file_type_lower == 'text/csv' or file_name.lower().endswith('.csv'):
				return FileContentExtractor._extract_from_csv(file_content)

			else:
				return None, f'Unsupported file type: {file_type}'

		except Exception as e:
			logger.error(f'Error extracting content from {file_name}: {str(e)}')
			return None, f'Extraction failed: {str(e)}'

	@staticmethod
	def _extract_from_txt(file_content: bytes) -> Tuple[Optional[str], Optional[str]]:
		"""Extract text from plain text files"""
		try:
			# Try different encodings
			for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
				try:
					text = file_content.decode(encoding)
					return text.strip(), None
				except UnicodeDecodeError:
					continue

			return None, 'Could not decode text file with any supported encoding'

		except Exception as e:
			return None, f'Error reading text file: {str(e)}'

	@staticmethod
	def _extract_from_pdf(file_content: bytes) -> Tuple[Optional[str], Optional[str]]:
		try:
			pdf_stream = io.BytesIO(file_content)
			pdf_reader = fitz.open(stream=pdf_stream, filetype='pdf')

			text_content = []
			for page_num in range(len(pdf_reader)):
				page = pdf_reader[page_num]
				text_content.append(page.get_text('text'))

			full_text = '\n'.join(text_content).strip()
			return full_text if full_text else None, None

		except ImportError:
			return None, 'PyMuPDF library not installed. Please install with: pip install PyMuPDF'
		except Exception as e:
			return None, f'Error reading PDF file: {str(e)}'

	@staticmethod
	def _extract_from_docx(file_content: bytes) -> Tuple[Optional[str], Optional[str]]:
		"""Extract text from DOCX files using python-docx"""
		try:
			from docx import Document

			# Create a temporary file to work with python-docx
			with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
				temp_file.write(file_content)
				temp_file.flush()

				try:
					doc = Document(temp_file.name)
					text_content = []

					for paragraph in doc.paragraphs:
						if paragraph.text.strip():
							text_content.append(paragraph.text)

					# Also extract text from tables
					for table in doc.tables:
						for row in table.rows:
							for cell in row.cells:
								if cell.text.strip():
									text_content.append(cell.text)

					full_text = '\n'.join(text_content).strip()
					return full_text if full_text else None, None

				finally:
					# Clean up temporary file
					try:
						os.unlink(temp_file.name)
					except:
						pass

		except ImportError:
			return None, 'python-docx library not installed. Please install with: pip install python-docx'
		except Exception as e:
			return None, f'Error reading DOCX file: {str(e)}'

	@staticmethod
	def _extract_from_doc(file_content: bytes) -> Tuple[Optional[str], Optional[str]]:
		"""Extract text from legacy DOC files"""
		try:
			# For legacy .doc files, we need python-docx2txt or similar
			# This is a basic implementation that may not work for all .doc files
			return None, 'Legacy .doc format not fully supported. Please convert to .docx format.'

		except Exception as e:
			return None, f'Error reading DOC file: {str(e)}'

	@staticmethod
	def _extract_from_csv(file_content: bytes) -> Tuple[Optional[str], Optional[str]]:
		"""Extract text from CSV files"""
		try:
			import csv

			# Try different encodings
			for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
				try:
					text_content = file_content.decode(encoding)

					# Parse CSV and convert to readable text
					csv_reader = csv.reader(io.StringIO(text_content))
					rows = []

					for row in csv_reader:
						if row:  # Skip empty rows
							rows.append(' | '.join(row))

					full_text = '\n'.join(rows).strip()
					return full_text if full_text else None, None

				except UnicodeDecodeError:
					continue

			return None, 'Could not decode CSV file with any supported encoding'

		except Exception as e:
			return None, f'Error reading CSV file: {str(e)}'

	@staticmethod
	def is_supported_file_type(file_type: str, file_name: str) -> bool:
		"""Check if a file type is supported for text extraction"""
		file_type_lower = file_type.lower()
		file_name_lower = file_name.lower()

		supported_types = ['text/plain', 'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml', 'application/msword', 'text/markdown', 'text/csv']

		supported_extensions = ['.txt', '.pdf', '.docx', '.doc', '.md', '.markdown', '.csv']

		# Check MIME type
		for supported_type in supported_types:
			if file_type_lower.startswith(supported_type):
				return True

		# Check file extension as fallback
		for ext in supported_extensions:
			if file_name_lower.endswith(ext):
				return True

		return False
